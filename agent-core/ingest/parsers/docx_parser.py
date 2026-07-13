from dataclasses import dataclass
from pathlib import Path

import docx


@dataclass
class DocxContent:
    paragraphs: list[str]
    tables: list[list[list[str]]]


def parse_docx(path: Path) -> DocxContent:
    """Parse a docx file into its paragraph text and any tables.

    Unlike parse_csv/parse_xlsx, tables here are NOT turned into
    header-keyed dicts -- a docx table isn't reliably a data export the
    way a CSV/xlsx sheet is, and assuming row 0 is a header would often
    be wrong. Cells stay as raw rows of strings; deciding whether a table
    has a header is an extraction-step judgment call, not a parsing one.
    Empty paragraphs (blank lines) are dropped -- they carry no content.
    """
    document = docx.Document(path)

    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]

    tables = [
        [[cell.text for cell in row.cells] for row in table.rows]
        for table in document.tables
    ]

    return DocxContent(paragraphs=paragraphs, tables=tables)
