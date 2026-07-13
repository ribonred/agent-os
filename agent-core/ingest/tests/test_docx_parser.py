from pathlib import Path

import docx

from parsers.docx_parser import parse_docx


def _write_docx(path: Path, paragraphs: list[str], table_rows: list[list[str]] | None = None) -> None:
    document = docx.Document()
    for text in paragraphs:
        document.add_paragraph(text)
    if table_rows:
        table = document.add_table(rows=0, cols=len(table_rows[0]))
        for row in table_rows:
            cells = table.add_row().cells
            for i, value in enumerate(row):
                cells[i].text = value
    document.save(path)


def test_parses_paragraphs(tmp_path: Path) -> None:
    docx_file = tmp_path / "intake.docx"
    _write_docx(docx_file, ["We are a solo skincare practice.", "Open Tue-Sat."])

    content = parse_docx(docx_file)

    assert content.paragraphs == [
        "We are a solo skincare practice.",
        "Open Tue-Sat.",
    ]
    assert content.tables == []


def test_drops_empty_paragraphs(tmp_path: Path) -> None:
    docx_file = tmp_path / "with_blanks.docx"
    _write_docx(docx_file, ["First line.", "", "   ", "Second line."])

    content = parse_docx(docx_file)

    assert content.paragraphs == ["First line.", "Second line."]


def test_parses_table_as_raw_rows_not_header_keyed(tmp_path: Path) -> None:
    docx_file = tmp_path / "pricelist.docx"
    _write_docx(
        docx_file,
        ["Price list:"],
        table_rows=[
            ["Service", "Price"],
            ["Facial", "45"],
            ["Consultation", "20"],
        ],
    )

    content = parse_docx(docx_file)

    assert content.tables == [
        [
            ["Service", "Price"],
            ["Facial", "45"],
            ["Consultation", "20"],
        ]
    ]
